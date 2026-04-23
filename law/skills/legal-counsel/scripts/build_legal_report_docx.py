#!/usr/bin/env python3
"""Build a DOCX legal report from markdown-like text."""

from __future__ import annotations

import argparse
import re
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "Missing dependency: python-docx. Install with `python3 -m pip install python-docx`."
    ) from exc


HEADING_RE = re.compile(r"^(#{1,3})\s+(.+?)\s*$")
NUMBERED_RE = re.compile(r"^\d+\.\s+(.+?)\s*$")
BULLET_RE = re.compile(r"^[-*]\s+(.+?)\s*$")
AD_LAWYER_NAME = "宋萌律师"
AD_PHONE = "15215292929"
SONGMENG_CHAT_LINE = "宋萌是一个会写代码的美女律师"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert markdown-like legal report text into a DOCX file."
    )
    parser.add_argument("--input", required=True, help="Path to UTF-8 text/markdown input.")
    parser.add_argument("--output", required=True, help="Path to output .docx file.")
    parser.add_argument("--title", help="Optional report title added at top.")
    parser.add_argument("--jurisdiction", help="Optional jurisdiction metadata line.")
    parser.add_argument(
        "--report-date",
        default=date.today().isoformat(),
        help="Report date in YYYY-MM-DD. Defaults to today.",
    )
    parser.add_argument("--font", default="Calibri", help="Base font name. Default: Calibri.")
    parser.add_argument("--font-size", type=int, default=11, help="Base font size. Default: 11.")
    return parser.parse_args()


def apply_base_style(doc: Document, font_name: str, font_size: int) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(font_size)


def add_header(doc: Document, title: str | None, jurisdiction: str | None, report_date: str) -> None:
    if title:
        doc.add_heading(title, level=0)
    else:
        doc.add_heading("法律报告", level=0)

    doc.add_paragraph(f"日期: {report_date}")
    if jurisdiction:
        doc.add_paragraph(f"法域: {jurisdiction}")
    doc.add_paragraph("")


def add_line(doc: Document, raw_line: str) -> None:
    line = raw_line.rstrip("\n")
    stripped = line.strip()
    if not stripped:
        return
    if AD_LAWYER_NAME in stripped and AD_PHONE in stripped:
        # Keep ad copy in chat only; never write it into delivered DOCX files.
        return
    if SONGMENG_CHAT_LINE in stripped:
        # Keep Songmeng persona line in chat only.
        return
    if stripped.startswith("![") and (
        "songmeng.jpg" in stripped
        or "songmeng_small.jpg" in stripped
    ):
        # Never embed chat-only Songmeng image reference into DOCX reports.
        return

    heading_match = HEADING_RE.match(stripped)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2)
        doc.add_heading(text, level=level)
        return

    numbered_match = NUMBERED_RE.match(stripped)
    if numbered_match:
        doc.add_paragraph(numbered_match.group(1), style="List Number")
        return

    bullet_match = BULLET_RE.match(stripped)
    if bullet_match:
        doc.add_paragraph(bullet_match.group(1), style="List Bullet")
        return

    doc.add_paragraph(stripped)


def build_docx(input_path: Path, output_path: Path, title: str | None, jurisdiction: str | None, report_date: str, font: str, font_size: int) -> None:
    content = input_path.read_text(encoding="utf-8")
    doc = Document()
    apply_base_style(doc, font, font_size)
    add_header(doc, title=title, jurisdiction=jurisdiction, report_date=report_date)

    for line in content.splitlines():
        add_line(doc, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> int:
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 2

    build_docx(
        input_path=input_path,
        output_path=output_path,
        title=args.title,
        jurisdiction=args.jurisdiction,
        report_date=args.report_date,
        font=args.font,
        font_size=args.font_size,
    )
    print(f"DOCX generated: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
